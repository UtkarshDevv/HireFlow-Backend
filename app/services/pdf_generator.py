"""
DOCX and PDF Resume Generator
─────────────────────────────
Generates professional .docx and .pdf resume files from structured content dicts.
Three templates:
  - clean_ats   : Simple, ATS-optimized, compact one-page layout
  - modern      : Section dividers, modern clean sans-serif typography
  - technical   : Compact one-page technical layout matching the supplied reference
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.config import get_settings

settings = get_settings()


def _ensure_dir() -> str:
    d = settings.resumes_dir
    os.makedirs(d, exist_ok=True)
    return d


def _add_horizontal_rule(doc: Document, color: str = "cbd5e1"):
    """Insert a thin horizontal line (bottom border on an empty paragraph)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")  # Thin line
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)


def _format_date_range(start: str, end: str, current: bool) -> str:
    end_str = "Present" if current else (end or "")
    if start and end_str:
        return f"{start} - {end_str}"
    return start or end_str or ""


# ─────────────────────────────────────────────────────────────────────────────
# DOCX Generation
# ─────────────────────────────────────────────────────────────────────────────

async def generate_docx(resume_id: str, content: dict, template_id: str = "clean_ats") -> str:
    """Generate a professional .docx file for the given resume content."""
    doc = Document()

    # Compact one-page A4 layout matching the reference resume.
    for section in doc.sections:
        section.top_margin = Inches(0.40)
        section.bottom_margin = Inches(0.40)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    # Compact ATS-friendly typography.
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9.5)
    font.color.rgb = RGBColor(0x1f, 0x1f, 0x1f)

    if template_id == "technical":
        _render_technical(doc, content)
    elif template_id == "modern":
        _render_modern(doc, content)
    else:
        _render_clean_ats(doc, content)

    out_dir = _ensure_dir()
    file_path = os.path.join(out_dir, f"{resume_id}.docx")
    doc.save(file_path)
    return file_path


def _ats_section(doc: Document, title: str, color_hex: str = "1e293b"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    _add_horizontal_rule(doc, color_hex)


def _render_clean_ats(doc: Document, content: dict):
    contact = content.get("contact", {})

    # Centered Header
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(contact.get("name", "Your Name"))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)  # Slate-900

    contact_parts = [
        contact.get("email", ""),
        contact.get("phone", ""),
        contact.get("location", ""),
        contact.get("linkedin", ""),
        contact.get("github", ""),
    ]
    contact_str = "  |  ".join(p for p in contact_parts if p)
    contact_para = doc.add_paragraph(contact_str)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(12)
    for run in contact_para.runs:
        run.font.size = Pt(9)

    _render_sections_shared(doc, content, "1e293b")


def _render_modern(doc: Document, content: dict):
    contact = content.get("contact", {})

    # Left-aligned Header
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(contact.get("name", "Your Name"))
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)  # Indigo-600

    contact_parts = [
        contact.get("email", ""),
        contact.get("phone", ""),
        contact.get("location", ""),
        contact.get("linkedin", ""),
        contact.get("github", ""),
    ]
    contact_str = "  ·  ".join(p for p in contact_parts if p)
    contact_para = doc.add_paragraph(contact_str)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact_para.paragraph_format.space_after = Pt(12)
    for run in contact_para.runs:
        run.font.size = Pt(9)

    _render_sections_shared(doc, content, "4f46e5")


def _render_technical(doc: Document, content: dict):
    """Compact technical resume matching the supplied one-page reference design."""
    contact = content.get("contact", {})

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(1)
    name_run = name_para.add_run(contact.get("name", "Your Name").upper())
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    line1 = " | ".join(
        x for x in [contact.get("location", ""), contact.get("phone", "")]
        if x
    )
    line2 = " | ".join(
        x for x in [
            contact.get("email", ""),
            contact.get("linkedin", ""),
            contact.get("github", ""),
        ] if x
    )

    for value in [line1, line2]:
        if value:
            p = doc.add_paragraph(value)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.font.size = Pt(8.5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(1)

    _render_sections_shared(doc, content, "111111", skip_skills=False)


def _render_sections_shared(doc: Document, content: dict, theme_color: str, skip_skills: bool = False):
    """Compact one-page section order:
    Summary -> Skills -> Experience -> Projects -> Education -> Certifications.
    """

    def add_compact_bullets(items, font_size=9.0):
        for bullet in items:
            if not bullet:
                continue
            bp = doc.add_paragraph(bullet, style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.20)
            bp.paragraph_format.first_line_indent = Inches(-0.12)
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(1)
            bp.paragraph_format.line_spacing = 1.0
            for run in bp.runs:
                run.font.size = Pt(font_size)

    # Professional Summary
    if content.get("summary"):
        _ats_section(doc, "PROFESSIONAL SUMMARY", theme_color)
        p = doc.add_paragraph(content["summary"])
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        for run in p.runs:
            run.font.size = Pt(9)

    # Technical Skills
    if not skip_skills and content.get("skills"):
        _ats_section(doc, "TECHNICAL SKILLS", theme_color)
        by_cat = {}
        for skill in content["skills"]:
            cat = skill.get("category", "General")
            name = skill.get("name", "")
            if name:
                by_cat.setdefault(cat, [])
                if name not in by_cat[cat]:
                    by_cat[cat].append(name)

        for cat, names in by_cat.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            run_cat = p.add_run(f"{cat}: ")
            run_cat.bold = True
            run_cat.font.size = Pt(9)
            p.add_run(", ".join(names)).font.size = Pt(9)

    # Experience
    if content.get("experience"):
        _ats_section(doc, "EXPERIENCE", theme_color)
        for exp in content["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.line_spacing = 1.0

            run_title = p.add_run(
                f"{exp.get('title', '')} | {exp.get('company', '')}"
            )
            run_title.bold = True
            run_title.font.size = Pt(9.5)

            date_str = _format_date_range(
                exp.get("start_date", ""),
                exp.get("end_date", ""),
                exp.get("current", False),
            )
            if date_str:
                run_date = p.add_run(f" | {date_str}")
                run_date.italic = True
                run_date.font.size = Pt(9)

            if exp.get("location"):
                run_loc = p.add_run(f" | {exp['location']}")
                run_loc.font.size = Pt(9)

            add_compact_bullets(exp.get("bullets", []), 9)

    # Projects
    if content.get("projects"):
        _ats_section(doc, "PROJECTS", theme_color)
        for proj in content["projects"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.line_spacing = 1.0

            run_proj = p.add_run(proj.get("name", ""))
            run_proj.bold = True
            run_proj.font.size = Pt(9.5)

            tech = proj.get("tech_stack", [])
            if tech:
                run_tech = p.add_run(f" | {', '.join(tech)}")
                run_tech.font.size = Pt(8.5)

            links = []
            if proj.get("github_url"):
                links.append("GitHub")
            if proj.get("url"):
                links.append("Live")
            if links:
                run_url = p.add_run(f" | {' · '.join(links)}")
                run_url.italic = True
                run_url.font.size = Pt(8.5)

            add_compact_bullets(proj.get("bullets", []), 9)

    # Education
    if content.get("education"):
        _ats_section(doc, "EDUCATION", theme_color)
        for edu in content["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

            school = edu.get("school", "")
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            date_str = _format_date_range(
                edu.get("start_date", ""),
                edu.get("end_date", ""),
                False,
            )

            run_school = p.add_run(school)
            run_school.bold = True
            run_school.font.size = Pt(9.5)

            detail = " — ".join(x for x in [degree + (f" in {field}" if field else "")] if x)
            if detail:
                run_detail = p.add_run(f" | {detail}")
                run_detail.font.size = Pt(9)

            if date_str:
                run_date = p.add_run(f" | {date_str}")
                run_date.italic = True
                run_date.font.size = Pt(8.5)

            if edu.get("gpa"):
                run_gpa = p.add_run(f" | GPA: {edu['gpa']}")
                run_gpa.font.size = Pt(8.5)

    # Certifications
    if content.get("certifications"):
        _ats_section(doc, "CERTIFICATIONS", theme_color)
        for cert in content["certifications"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            run_cert = p.add_run(cert.get("name", ""))
            run_cert.bold = True
            run_cert.font.size = Pt(9)
            info = " | ".join(
                x for x in [cert.get("issuer", ""), cert.get("date", "")]
                if x
            )
            if info:
                p.add_run(f" | {info}").font.size = Pt(8.5)


# ─────────────────────────────────────────────────────────────────────────────
# PDF Generation using fpdf2
# ─────────────────────────────────────────────────────────────────────────────

def _clean_for_pdf(data):
    """Recursively sanitize dictionary data to ensure latin-1 compatibility for fpdf2."""
    if isinstance(data, dict):
        return {k: _clean_for_pdf(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [_clean_for_pdf(x) for x in data]
    elif isinstance(data, str):
        replacements = {
            "\u2013": "-",
            "\u2014": "-",
            "\u2018": "'",
            "\u2019": "'",
            "\u201b": "'",
            "\u201c": '"',
            "\u201d": '"',
            "\u201f": '"',
            "\u2022": "-",
            "\u00a0": " ",
            "—": "-",
            "–": "-",
            "·": "|",
        }
        for orig, repl in replacements.items():
            data = data.replace(orig, repl)
        return data.encode("latin-1", errors="ignore").decode("latin-1")
    return data


def _clean_summary(summary_text: str) -> str:
    """
    If the summary string contains copy-pasted resume sections 
    (like 'EXPERIENCE', 'TECHNICAL SKILLS', 'EDUCATION', etc.),
    truncate it to retain ONLY the real summary paragraph.
    """
    if not summary_text:
        return ""
    
    keywords = [
        "EXPERIENCE", "TECHNICAL SKILLS", "EDUCATION", "PROJECTS", 
        "WORK EXPERIENCE", "SKILLS", "CERTIFICATIONS", "LANGUAGES:"
    ]
    
    lines = summary_text.split("\n")
    clean_lines = []
    for line in lines:
        stripped = line.strip()
        if any(stripped == kw or stripped.startswith(kw + " ") for kw in keywords):
            break
        inline_found = False
        for kw in keywords:
            if f" {kw}" in line or f"\t{kw}" in line:
                line = line.split(f" {kw}")[0].split(f"\t{kw}")[0]
                inline_found = True
                break
        clean_lines.append(line)
        if inline_found:
            break
            
    return "\n".join(clean_lines).strip()


async def generate_pdf(resume_id: str, content: dict, template_id: str = "clean_ats", tailoring_meta: dict = None) -> str:
    """Generate a beautiful, professional PDF file for the given resume content based on LaTeX SDE template."""
    from fpdf import FPDF
    
    # Sanitize content for PDF core fonts (latin-1)
    content = _clean_for_pdf(content)
    
    # Setup PDF document
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    
    # Compact A4 margins matching the supplied one-page reference.
    left_margin = 14
    top_margin = 10
    right_margin = 14
    printable_width = 210 - left_margin - right_margin  # 172mm
    
    pdf.set_margins(left=left_margin, top=top_margin, right=right_margin)
    
    font_family = "Helvetica"
    
    # Setup Colors matching LaTeX dark text + black rules
    text_color = (0, 0, 0)         # Crisp black
    body_color = (25, 25, 25)       # Off-black for readability
    rule_color = (0, 0, 0)         # Solid section rule
    
    contact = content.get("contact", {})

    # --- HEADER: centered, compact, reference-style ---
    pdf.set_font(font_family, "B", 18)
    pdf.set_text_color(*text_color)
    pdf.cell(printable_width, 7, contact.get("name", "Your Name").upper(), ln=1, align="C")

    pdf.set_font(font_family, "", 8.5)
    pdf.set_text_color(*body_color)

    line1 = " | ".join(
        x for x in [contact.get("location", ""), contact.get("phone", "")]
        if x
    )
    if line1:
        pdf.cell(printable_width, 4, line1, ln=1, align="C")

    # Header Contact Links
    header_links = []
    if contact.get("email"):
        header_links.append(("Email", f"mailto:{contact['email']}", contact['email']))
    if contact.get("linkedin"):
        ln_url = contact['linkedin'] if contact['linkedin'].startswith("http") else f"https://{contact['linkedin']}"
        header_links.append(("LinkedIn", ln_url, contact['linkedin']))
    if contact.get("github"):
        gh_url = contact['github'] if contact['github'].startswith("http") else f"https://{contact['github']}"
        header_links.append(("GitHub", gh_url, contact['github']))
    if contact.get("website"):
        ws_url = contact['website'] if contact['website'].startswith("http") else f"https://{contact['website']}"
        header_links.append(("Website", ws_url, contact['website']))

    if header_links:
        # Re-center links
        pdf.set_x(left_margin)
        link_str = " | ".join([val for _, _, val in header_links])
        # A simple estimation of centering by calculating width is complex in fpdf directly. 
        # Here we just output centered using dummy cell.
        pdf.cell(printable_width, 4, "", ln=1) 
        pdf.set_y(pdf.get_y() - 4)
        
        # Actually print the links as individual writes for functionality
        full_line = " | ".join([val for _, _, val in header_links])
        total_w = pdf.get_string_width(full_line)
        pdf.set_x((210 - total_w) / 2)
        for idx, (label, href, display_val) in enumerate(header_links):
            if idx > 0:
                pdf.write(4, " | ")
            pdf.set_font(font_family, "U", 8.5)
            pdf.set_text_color(37, 99, 235)
            pdf.write(4, display_val, link=href)
        pdf.ln(4.5)

    def add_section(title_text):
        pdf.ln(2)
        pdf.set_font(font_family, "B", 10)
        pdf.set_text_color(*text_color)
        pdf.cell(printable_width, 4.5, title_text.upper(), ln=1)

        pdf.set_draw_color(*rule_color)
        pdf.set_line_width(0.25)
        pdf.line(left_margin, pdf.get_y(), 210 - right_margin, pdf.get_y())
        pdf.ln(1.5)

    def add_bullets(items):
        for bullet in items:
            if not bullet:
                continue
            pdf.set_font(font_family, "", 8.5)
            pdf.set_text_color(*body_color)
            pdf.set_x(left_margin + 4)
            pdf.cell(3, 3.6, "-", ln=0)
            pdf.set_x(left_margin + 7)
            pdf.multi_cell(printable_width - 7, 3.6, bullet, ln=1)

    # --- PROFESSIONAL SUMMARY ---
    clean_sum = _clean_summary(content.get("summary", ""))
    if clean_sum:
        add_section("PROFESSIONAL SUMMARY")
        pdf.set_font(font_family, "", 8.5)
        pdf.multi_cell(printable_width, 3.8, clean_sum, ln=1)

    # --- TECHNICAL SKILLS ---
    if content.get("skills"):
        add_section("TECHNICAL SKILLS")
        by_cat = {}
        for skill in content["skills"]:
            cat = skill.get("category", "General").strip()
            name = skill.get("name", "").strip()
            if name:
                by_cat.setdefault(cat, [])
                if name not in by_cat[cat]:
                    by_cat[cat].append(name)

        for cat, names in by_cat.items():
            pdf.set_font(font_family, "B", 8.5)
            pdf.set_text_color(*text_color)
            label = f"{cat}: "
            label_w = 31
            pdf.cell(label_w, 3.8, label, ln=0)
            pdf.set_font(font_family, "", 8.5)
            pdf.set_text_color(*body_color)
            pdf.multi_cell(printable_width - label_w, 3.8, ", ".join(names), ln=1)

    # --- EXPERIENCE ---
    if content.get("experience"):
        add_section("EXPERIENCE")
        seen_exp = set()
        unique_exp = []
        for exp in content["experience"]:
            key = (
                exp.get("company", "").strip().lower(),
                exp.get("title", "").strip().lower(),
            )
            if key not in seen_exp:
                seen_exp.add(key)
                unique_exp.append(exp)

        for exp in unique_exp:
            pdf.set_font(font_family, "B", 8.8)
            pdf.set_text_color(*text_color)

            role = exp.get("title", "")
            company = exp.get("company", "")
            date_str = _format_date_range(
                exp.get("start_date", ""),
                exp.get("end_date", ""),
                exp.get("current", False),
            )

            title_line = " | ".join(x for x in [role, company] if x)
            if date_str:
                title_line += f" | {date_str}"
            if exp.get("location"):
                title_line += f" | {exp['location']}"

            pdf.cell(printable_width, 4, title_line, ln=1)
            add_bullets(exp.get("bullets", []))

    # --- PROJECTS ---
    if content.get("projects"):
        add_section("PROJECTS")
        for proj in content["projects"]:
            p_name = proj.get("name", "")
            tech_stack = proj.get("tech_stack") or []
            gh_raw = proj.get("github_url") or proj.get("github") or proj.get("github_link") or ""
            proj_url_raw = proj.get("url") or proj.get("project_url") or proj.get("link") or proj.get("live_url") or ""

            # Project title
            pdf.set_font(font_family, "B", 8.8)
            pdf.set_text_color(*text_color)
            pdf.write(4, p_name)

            if tech_stack:
                pdf.set_font(font_family, "", 8.5)
                pdf.set_text_color(80, 80, 80)
                pdf.write(4, f" | {', '.join(tech_stack)}")

            if gh_raw:
                pdf.set_font(font_family, "U", 8.5)
                pdf.set_text_color(37, 99, 235)  # blue link
                gh_href = gh_raw if gh_raw.startswith("http") else f"https://{gh_raw}"
                pdf.write(4, " | [GitHub]", link=gh_href)

            if proj_url_raw:
                pdf.set_font(font_family, "U", 8.5)
                pdf.set_text_color(37, 99, 235)  # blue link
                url_href = proj_url_raw if proj_url_raw.startswith("http") else f"https://{proj_url_raw}"
                pdf.write(4, " | [Live Demo]", link=url_href)

            pdf.ln(4.2)
            add_bullets(proj.get("bullets", []))

    # --- EDUCATION ---
    if content.get("education"):
        add_section("EDUCATION")
        for edu in content["education"]:
            pdf.set_font(font_family, "B", 8.8)
            pdf.set_text_color(*text_color)

            school = edu.get("school", "")
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            date_str = _format_date_range(
                edu.get("start_date", ""),
                edu.get("end_date", ""),
                False,
            )

            line = school
            details = degree + (f" in {field}" if field else "")
            if details:
                line += f" | {details}"
            if date_str:
                line += f" | {date_str}"
            if edu.get("gpa"):
                line += f" | GPA: {edu['gpa']}"

            pdf.multi_cell(printable_width, 4, line, ln=1)

    # --- CERTIFICATIONS ---
    if content.get("certifications"):
        add_section("CERTIFICATIONS")
        for cert in content["certifications"]:
            pdf.set_font(font_family, "B", 8.5)
            pdf.set_text_color(*text_color)
            name = cert.get("name", "")
            info = " | ".join(
                x for x in [cert.get("issuer", ""), cert.get("date", "")]
                if x
            )
            pdf.cell(printable_width, 3.8, name + (f" | {info}" if info else ""), ln=1)

    out_dir = _ensure_dir()
    file_path = os.path.join(out_dir, f"{resume_id}.pdf")
    pdf.output(file_path)
    return file_path